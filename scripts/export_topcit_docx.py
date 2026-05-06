#!/usr/bin/env python3
"""Export TOPCIT practice questions and answers to separate DOCX files."""

from __future__ import annotations

import json
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
DATA_PATH = ROOT / "topcit_json_v2" / "topcit_practice.json"
OUTPUT_DIR = ROOT / "output" / "doc"
QUESTIONS_PATH = OUTPUT_DIR / "topcit_1050_questions.docx"
ANSWERS_PATH = OUTPUT_DIR / "topcit_1050_answers.docx"

CHAPTER_TITLES = {
    "01": "소프트웨어 개발",
    "02": "데이터 이해와 활용",
    "03": "시스템 아키텍처 이해와 활용",
    "04": "정보보안 이해와 활용",
    "05": "IT 비즈니스와 윤리",
    "06": "테크니컬 커뮤니케이션과 프로젝트 관리",
}


def set_run_font(run, name: str = "맑은 고딕", size_pt: float | None = None) -> None:
    run.font.name = name
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def set_paragraph_spacing(paragraph, before: float = 0, after: float = 3, line: float = 1.08) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "맑은 고딕"
    normal.font.size = Pt(9.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.space_after = Pt(3)

    for style_name, size in (("Title", 18), ("Heading 1", 14), ("Heading 2", 11)):
        style = styles[style_name]
        style.font.name = "맑은 고딕"
        style.font.size = Pt(size)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, size_pt=8)


def add_title(document: Document, title: str, subtitle: str) -> None:
    p = document.add_paragraph()
    p.style = document.styles["Title"]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    set_run_font(r, size_pt=18)
    r.bold = True

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    set_run_font(r, size_pt=9)
    set_paragraph_spacing(p, after=10)


def add_question_block(document: Document, index: int, item: dict) -> None:
    p = document.add_paragraph()
    set_paragraph_spacing(p, before=3, after=2)
    meta = f"{index}. [{item['id']}] {item['type']} | {item['domain']} | {item['topic']}"
    r = p.add_run(meta)
    set_run_font(r, size_pt=9.5)
    r.bold = True

    p = document.add_paragraph()
    set_paragraph_spacing(p, after=5)
    for line_index, line in enumerate(item["question"].splitlines()):
        if line_index:
            p.add_run().add_break()
        r = p.add_run(line)
        set_run_font(r, size_pt=9.5)

    if item.get("choices"):
        for choice in item["choices"]:
            p = document.add_paragraph(style=None)
            p.paragraph_format.left_indent = Cm(0.35)
            set_paragraph_spacing(p, after=1)
            r = p.add_run(str(choice))
            set_run_font(r, size_pt=9.2)

    spacer = document.add_paragraph()
    set_paragraph_spacing(spacer, before=0, after=4, line=1.0)


def add_answer_block(document: Document, index: int, item: dict) -> None:
    p = document.add_paragraph()
    set_paragraph_spacing(p, before=3, after=1)
    r = p.add_run(f"{index}. [{item['id']}] {item['topic']}")
    set_run_font(r, size_pt=9.5)
    r.bold = True

    p = document.add_paragraph()
    set_paragraph_spacing(p, after=1)
    r = p.add_run(f"정답: {item['answer']}")
    set_run_font(r, size_pt=9.5)
    r.bold = True

    explanation = item.get("explanation", "").strip()
    if explanation:
        p = document.add_paragraph()
        set_paragraph_spacing(p, after=5)
        r = p.add_run(f"해설: {explanation}")
        set_run_font(r, size_pt=9)


def grouped_questions(questions: list[dict]) -> dict[str, list[dict]]:
    groups: dict[str, list[dict]] = defaultdict(list)
    for item in questions:
        groups[item["id"].split("-")[1]].append(item)
    return dict(sorted(groups.items()))


def add_chapter_heading(document: Document, chapter_id: str, count: int) -> None:
    if len(document.paragraphs) > 2:
        document.add_section(WD_SECTION.NEW_PAGE)
    title = CHAPTER_TITLES.get(chapter_id, f"{chapter_id}과목")
    heading = document.add_heading(f"{chapter_id}. {title} ({count}문항)", level=1)
    set_paragraph_spacing(heading, before=0, after=8)


def build_questions_doc(questions: list[dict]) -> Document:
    document = Document()
    configure_document(document)
    add_title(
        document,
        "TOPCIT 예상문제 1050 - 문제",
        "정답과 해설은 별도 문서(topcit_1050_answers.docx)에 분리되어 있습니다.",
    )
    groups = grouped_questions(questions)
    number = 1
    for chapter_id, items in groups.items():
        add_chapter_heading(document, chapter_id, len(items))
        for item in items:
            add_question_block(document, number, item)
            number += 1
    add_page_number(document.sections[0].footer.paragraphs[0])
    return document


def build_answers_doc(questions: list[dict]) -> Document:
    document = Document()
    configure_document(document)
    add_title(
        document,
        "TOPCIT 예상문제 1050 - 정답",
        "문제 문서의 번호와 문항 ID를 기준으로 대조하세요.",
    )
    groups = grouped_questions(questions)
    number = 1
    for chapter_id, items in groups.items():
        add_chapter_heading(document, chapter_id, len(items))
        for item in items:
            add_answer_block(document, number, item)
            number += 1
    add_page_number(document.sections[0].footer.paragraphs[0])
    return document


def main() -> None:
    with DATA_PATH.open(encoding="utf-8") as f:
        data = json.load(f)

    questions = data["questions"]
    if len(questions) != 1050:
        raise RuntimeError(f"Expected 1050 questions, found {len(questions)}")

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    build_questions_doc(questions).save(QUESTIONS_PATH)
    build_answers_doc(questions).save(ANSWERS_PATH)
    print(QUESTIONS_PATH)
    print(ANSWERS_PATH)


if __name__ == "__main__":
    main()
